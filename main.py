#!/usr/bin/env python3
"""
Booksy Desktop - Professional Book Creator
A standalone desktop application for creating and managing books
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import json
import os
from datetime import datetime
import uuid
import threading
from pathlib import Path

class BooksyDesktop:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Booksy Desktop - Write. Create. Publish.")
        self.root.geometry("1200x800")
        self.root.configure(bg='#e6ebe0')
        
        # Data storage
        self.data_dir = Path.home() / "Booksy"
        self.data_dir.mkdir(exist_ok=True)
        self.books_file = self.data_dir / "books.json"
        
        # Current state
        self.current_book = None
        self.current_section = None
        self.books = self.load_books()
        
        # Colors (Booksy palette)
        self.colors = {
            'primary': '#36c9c6',
            'secondary': '#9bc1bc',
            'accent': '#ed6a5a',
            'light_bg': '#e6ebe0',
            'warm_bg': '#f4f1bb',
            'dark_text': '#2c3e50'
        }
        
        self.setup_ui()
        self.show_dashboard()
        
    def setup_ui(self):
        # Main container
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Configure styles
        style = ttk.Style()
        style.configure('Title.TLabel', font=('Arial', 16, 'bold'))
        style.configure('Header.TLabel', font=('Arial', 12, 'bold'))
        
    def show_dashboard(self):
        # Clear main frame
        for widget in self.main_frame.winfo_children():
            widget.destroy()
            
        # Header
        header_frame = ttk.Frame(self.main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        ttk.Label(header_frame, text="📚 My Books", style='Title.TLabel').pack(side=tk.LEFT)
        ttk.Button(header_frame, text="+ New Book", command=self.show_create_book).pack(side=tk.RIGHT)
        
        # Books list
        if not self.books:
            self.show_empty_state()
        else:
            self.show_books_list()
    
    def show_empty_state(self):
        empty_frame = ttk.Frame(self.main_frame)
        empty_frame.pack(expand=True, fill=tk.BOTH)
        
        ttk.Label(empty_frame, text="📖", font=('Arial', 48)).pack(pady=20)
        ttk.Label(empty_frame, text="No books yet", font=('Arial', 18, 'bold')).pack()
        ttk.Label(empty_frame, text="Create your first book to get started").pack(pady=10)
        ttk.Button(empty_frame, text="Create First Book", command=self.show_create_book).pack(pady=20)
    
    def show_books_list(self):
        # Books frame with scrollbar
        canvas = tk.Canvas(self.main_frame, bg='#e6ebe0')
        scrollbar = ttk.Scrollbar(self.main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Create book cards
        for i, (book_id, book) in enumerate(self.books.items()):
            self.create_book_card(scrollable_frame, book_id, book, i)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
    
    def create_book_card(self, parent, book_id, book, index):
        # Card frame
        card = ttk.LabelFrame(parent, text=f"📚 {book['title']}", padding=15)
        card.pack(fill=tk.X, pady=10, padx=5)
        
        # Book info
        info_frame = ttk.Frame(card)
        info_frame.pack(fill=tk.X)
        
        ttk.Label(info_frame, text=f"by {book['author']}", font=('Arial', 10, 'italic')).pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"Format: {book['format'].title()}", font=('Arial', 9)).pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"Updated: {book['updated_at'][:10]}", font=('Arial', 9)).pack(anchor=tk.W)
        
        # Stats
        word_count = sum(len(content.split()) for content in book.get('content', {}).values() if content)
        ttk.Label(info_frame, text=f"Words: {word_count}", font=('Arial', 9, 'bold')).pack(anchor=tk.W)
        
        # Buttons
        btn_frame = ttk.Frame(card)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(btn_frame, text="✏️ Edit", command=lambda: self.edit_book(book_id)).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="📄 Export DOCX", command=lambda: self.export_book(book_id, 'docx')).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🗑️ Delete", command=lambda: self.delete_book(book_id)).pack(side=tk.RIGHT)
    
    def show_create_book(self):
        # Clear main frame
        for widget in self.main_frame.winfo_children():
            widget.destroy()
            
        # Header
        header_frame = ttk.Frame(self.main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        ttk.Label(header_frame, text="📝 Create New Book", style='Title.TLabel').pack(side=tk.LEFT)
        ttk.Button(header_frame, text="← Back", command=self.show_dashboard).pack(side=tk.RIGHT)
        
        # Form
        form_frame = ttk.Frame(self.main_frame)
        form_frame.pack(pady=20)
        
        ttk.Label(form_frame, text="Book Title:", style='Header.TLabel').grid(row=0, column=0, sticky=tk.W, pady=5)
        self.title_entry = ttk.Entry(form_frame, width=40, font=('Arial', 12))
        self.title_entry.grid(row=0, column=1, pady=5, padx=(10, 0))
        
        ttk.Label(form_frame, text="Author:", style='Header.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        self.author_entry = ttk.Entry(form_frame, width=40, font=('Arial', 12))
        self.author_entry.grid(row=1, column=1, pady=5, padx=(10, 0))
        
        ttk.Label(form_frame, text="Format:", style='Header.TLabel').grid(row=2, column=0, sticky=tk.W, pady=5)
        self.format_var = tk.StringVar()
        format_combo = ttk.Combobox(form_frame, textvariable=self.format_var, width=37, font=('Arial', 12))
        format_combo['values'] = ('Novel', 'Poetry Collection', 'Memoir', 'Cookbook', 'Children\'s Book', 'Technical/Business')
        format_combo.grid(row=2, column=1, pady=5, padx=(10, 0))
        
        ttk.Button(form_frame, text="Create Book", command=self.create_book).grid(row=3, column=1, pady=20, sticky=tk.E)
    
    def create_book(self):
        title = self.title_entry.get().strip()
        author = self.author_entry.get().strip()
        format_name = self.format_var.get()
        
        if not all([title, author, format_name]):
            messagebox.showerror("Error", "Please fill in all fields")
            return
        
        book_id = str(uuid.uuid4())
        format_key = format_name.lower().replace(' ', '_').replace('/', '_')
        
        # Get format-specific content structure
        content = self.get_format_content(format_key, title, author)
        
        book_data = {
            'id': book_id,
            'title': title,
            'author': author,
            'format': format_key,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat(),
            'content': content
        }
        
        self.books[book_id] = book_data
        self.save_books()
        
        messagebox.showinfo("Success", f"Book '{title}' created successfully!")
        self.edit_book(book_id)
    
    def get_format_content(self, format_key, title, author):
        """Get format-specific content structure"""
        year = datetime.now().year
        
        base_content = {
            'title_page': f"{title}\n\nby {author}",
            'copyright': f"Copyright © {year} by {author}\nAll rights reserved.",
            'dedication': "[Dedicate your book to someone special...]",
            'acknowledgments': "[Thank those who helped make this book possible...]"
        }
        
        if format_key == 'novel':
            base_content.update({
                'prologue': "# Prologue\n\n[Set the stage for your story...]",
                'chapter_1': "# Chapter 1\n\n[Begin your story here...]",
                'chapter_2': "# Chapter 2\n\n[Continue your narrative...]",
                'epilogue': "# Epilogue\n\n[Conclude your story...]",
                'about_author': f"About the Author\n\n{author} is..."
            })
        
        elif format_key == 'poetry_collection':
            base_content.update({
                'introduction': "# Introduction\n\n[Introduce your poetry collection...]",
                'section_1_love': "# Love & Relationships\n\n[Your first poem here...]\n\n---\n\n[Second poem here...]",
                'section_2_nature': "# Nature & Seasons\n\n[Nature-themed poems...]",
                'section_3_life': "# Life & Growth\n\n[Reflective poems...]",
                'section_4_dreams': "# Dreams & Aspirations\n\n[Inspirational poems...]",
                'notes': "# Notes\n\n[Background on your poems...]"
            })
        
        elif format_key == 'memoir':
            base_content.update({
                'foreword': "# Foreword\n\n[Why you're telling your story...]",
                'chapter_1_early_years': "# Chapter 1: Early Years\n\n[Your childhood memories...]",
                'chapter_2_growing_up': "# Chapter 2: Growing Up\n\n[Your teenage years...]",
                'afterword': "# Afterword\n\n[Final thoughts and hopes...]"
            })
        
        elif format_key == 'cookbook':
            base_content.update({
                'introduction': "# Introduction\n\n[Your cooking philosophy and story...]",
                'appetizers': "# Appetizers & Starters\n\n## Recipe Name\n\n**Ingredients:**\n- Ingredient 1\n- Ingredient 2\n\n**Instructions:**\n1. Step 1\n2. Step 2\n\n**Serves:** 4\n**Prep Time:** 15 minutes\n**Cook Time:** 20 minutes",
                'main_courses': "# Main Courses\n\n[Your main dish recipes...]",
                'desserts': "# Desserts\n\n[Your dessert recipes...]"
            })
        
        elif format_key == 'children_s_book':
            base_content.update({
                'chapter_1': "# Chapter 1\n\n[Once upon a time...]\n\n[Illustration note: Describe the scene for an illustrator]",
                'chapter_2': "# Chapter 2\n\n[Continue the adventure...]\n\n[Illustration note: What should be shown here]",
                'activities': "# Fun Activities\n\n[Coloring pages, puzzles, or games related to your story...]"
            })
        
        elif format_key == 'technical_business':
            base_content.update({
                'preface': "# Preface\n\n[Why this book is needed...]",
                'chapter_1_introduction': "# Chapter 1: Introduction\n\n[Introduce the topic and objectives...]",
                'chapter_2_fundamentals': "# Chapter 2: Fundamentals\n\n[Basic concepts and principles...]",
                'conclusion': "# Conclusion\n\n[Summary and final thoughts...]",
                'bibliography': "# Bibliography\n\n[Sources and references...]"
            })
        
        return base_content
    
    def edit_book(self, book_id):
        self.current_book = book_id
        book = self.books[book_id]
        
        # Clear main frame
        for widget in self.main_frame.winfo_children():
            widget.destroy()
        
        # Create editor layout
        self.setup_editor(book)
    
    def setup_editor(self, book):
        # Header
        header_frame = ttk.Frame(self.main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(header_frame, text=f"✏️ {book['title']}", style='Title.TLabel').pack(side=tk.LEFT)
        
        btn_frame = ttk.Frame(header_frame)
        btn_frame.pack(side=tk.RIGHT)
        
        ttk.Button(btn_frame, text="← Dashboard", command=self.show_dashboard).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="💾 Save", command=self.save_current_content).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="📄 Export", command=lambda: self.export_book(self.current_book, 'docx')).pack(side=tk.LEFT, padx=5)
        
        # Main editor area
        editor_frame = ttk.Frame(self.main_frame)
        editor_frame.pack(fill=tk.BOTH, expand=True)
        
        # Sidebar for sections
        sidebar = ttk.LabelFrame(editor_frame, text="Book Structure", width=200)
        sidebar.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        sidebar.pack_propagate(False)
        
        # Sections list
        self.sections_frame = ttk.Frame(sidebar)
        self.sections_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.update_sections_list()
        
        # Editor area
        editor_container = ttk.Frame(editor_frame)
        editor_container.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Current section label
        self.section_label = ttk.Label(editor_container, text="Select a section to edit", style='Header.TLabel')
        self.section_label.pack(pady=(0, 10))
        
        # Font controls
        font_frame = ttk.Frame(editor_container)
        font_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(font_frame, text="Font:").pack(side=tk.LEFT, padx=(0, 5))
        
        self.font_var = tk.StringVar(value="Georgia")
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_var, width=12, values=[
            "Georgia", "Arial", "Times New Roman", "Calibri", "Verdana", "Courier New"
        ])
        font_combo.pack(side=tk.LEFT, padx=(0, 10))
        font_combo.bind('<<ComboboxSelected>>', self.on_font_change)
        
        ttk.Label(font_frame, text="Size:").pack(side=tk.LEFT, padx=(0, 5))
        
        self.size_var = tk.StringVar(value="12")
        size_combo = ttk.Combobox(font_frame, textvariable=self.size_var, width=5, values=[
            "8", "9", "10", "11", "12", "14", "16", "18", "20", "24", "28", "32"
        ])
        size_combo.pack(side=tk.LEFT, padx=(0, 10))
        size_combo.bind('<<ComboboxSelected>>', self.on_size_change)
        
        # Initialize selection storage
        self.stored_selection = None
        
        # Text editor
        self.text_editor = scrolledtext.ScrolledText(
            editor_container,
            wrap=tk.WORD,
            font=('Georgia', 12),
            height=25,
            bg='white',
            fg='#2c3e50'
        )
        self.text_editor.pack(fill=tk.BOTH, expand=True)
        
        # Word count
        self.word_count_label = ttk.Label(editor_container, text="Words: 0")
        self.word_count_label.pack(pady=(5, 0))
        
        # Bind events
        self.text_editor.bind('<KeyRelease>', self.update_word_count)
        self.text_editor.bind('<Button-1>', self.update_word_count)
    
    def update_sections_list(self):
        for widget in self.sections_frame.winfo_children():
            widget.destroy()
        
        book = self.books[self.current_book]
        sections = book.get('content', {})
        
        # Define section order
        section_order = self.get_section_order(book['format'], sections)
        
        for section_key in section_order:
            if section_key not in sections:
                continue
                
            section_name = section_key.replace('_', ' ').title()
            
            # Create frame for section button and delete button
            section_frame = ttk.Frame(self.sections_frame)
            section_frame.pack(fill=tk.X, pady=2)
            
            # Section button
            btn = ttk.Button(
                section_frame,
                text=section_name,
                command=lambda s=section_key: self.load_section(s)
            )
            btn.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            # Delete button for chapters beyond first 2
            if section_key.startswith('chapter_') and section_key not in ['chapter_1', 'chapter_2']:
                del_btn = ttk.Button(
                    section_frame,
                    text="🗑️",
                    width=3,
                    command=lambda s=section_key: self.delete_chapter(s)
                )
                del_btn.pack(side=tk.RIGHT, padx=(2, 0))
            
            # Add chapter button after last chapter for applicable formats
            if (book['format'] in ['novel', 'memoir', 'technical_business'] and 
                section_key.startswith('chapter_') and 
                self.is_last_chapter(section_key, sections)):
                ttk.Button(
                    self.sections_frame,
                    text="+ Add Chapter",
                    command=self.add_chapter
                ).pack(fill=tk.X, pady=5)
    
    def get_section_order(self, format_key, sections):
        """Get the correct order of sections for display"""
        base_order = ['title_page', 'copyright', 'dedication', 'acknowledgments']
        
        if format_key == 'novel':
            base_order.extend(['prologue'])
            # Add chapters in order
            chapters = sorted([k for k in sections.keys() if k.startswith('chapter_')], 
                            key=lambda x: int(x.split('_')[1]))
            base_order.extend(chapters)
            base_order.extend(['epilogue', 'about_author'])
        
        elif format_key == 'memoir':
            base_order.extend(['foreword'])
            chapters = sorted([k for k in sections.keys() if k.startswith('chapter_')], 
                            key=lambda x: int(x.split('_')[1]) if x.split('_')[1].isdigit() else 999)
            base_order.extend(chapters)
            base_order.extend(['afterword'])
        
        elif format_key == 'technical_business':
            base_order.extend(['preface'])
            chapters = sorted([k for k in sections.keys() if k.startswith('chapter_')], 
                            key=lambda x: int(x.split('_')[1]) if x.split('_')[1].isdigit() else 999)
            base_order.extend(chapters)
            base_order.extend(['conclusion', 'bibliography'])
        
        else:
            # For other formats, just add remaining sections
            remaining = [k for k in sections.keys() if k not in base_order]
            base_order.extend(remaining)
        
        return base_order
    
    def is_last_chapter(self, chapter_key, sections):
        """Check if this is the last chapter"""
        chapters = [k for k in sections.keys() if k.startswith('chapter_')]
        if not chapters:
            return False
        last_chapter = max(chapters, key=lambda x: int(x.split('_')[1]))
        return chapter_key == last_chapter
    
    def load_section(self, section_key):
        self.current_section = section_key
        book = self.books[self.current_book]
        content = book.get('content', {}).get(section_key, '')
        
        self.section_label.config(text=f"Editing: {section_key.replace('_', ' ').title()}")
        
        self.text_editor.delete(1.0, tk.END)
        self.text_editor.insert(1.0, content)
        
        self.update_word_count()
    
    def add_chapter(self):
        book = self.books[self.current_book]
        chapters = [k for k in book.get('content', {}).keys() if k.startswith('chapter_')]
        next_num = len(chapters) + 1
        
        chapter_key = f'chapter_{next_num}'
        book['content'][chapter_key] = f"# Chapter {next_num}\n\n[Write your chapter content here...]"
        
        self.save_books()
        self.update_sections_list()
        self.load_section(chapter_key)
    
    def delete_chapter(self, chapter_key):
        book = self.books[self.current_book]
        chapter_name = chapter_key.replace('_', ' ').title()
        
        if messagebox.askyesno("Delete Chapter", f"Delete '{chapter_name}'?\nThis cannot be undone."):
            del book['content'][chapter_key]
            book['updated_at'] = datetime.now().isoformat()
            self.save_books()
            
            # Clear editor if deleted chapter was selected
            if self.current_section == chapter_key:
                self.current_section = None
                self.text_editor.delete(1.0, tk.END)
                self.section_label.config(text="Select a section to edit")
            
            self.update_sections_list()
    

    
    def on_font_change(self, event=None):
        self.store_selection()
        self.apply_stored_formatting()
    
    def on_size_change(self, event=None):
        self.store_selection()
        self.apply_stored_formatting()
    
    def store_selection(self):
        try:
            sel_start = self.text_editor.index(tk.SEL_FIRST)
            sel_end = self.text_editor.index(tk.SEL_LAST)
            self.stored_selection = (sel_start, sel_end)
        except tk.TclError:
            # Try to use stored selection if available
            pass
    
    def apply_stored_formatting(self):
        if self.stored_selection:
            sel_start, sel_end = self.stored_selection
            font_family = self.font_var.get()
            font_size = int(self.size_var.get())
            
            # Create a unique tag for the selected text
            import time
            tag_name = f"font_{font_family}_{font_size}_{int(time.time() * 1000)}"
            self.text_editor.tag_add(tag_name, sel_start, sel_end)
            self.text_editor.tag_config(tag_name, font=(font_family, font_size))
            
            # Clear stored selection after applying
            self.stored_selection = None
        else:
            messagebox.showwarning("No Selection", "Please select text first, then choose font/size.")
    
    def apply_font_to_selection(self):
        self.store_selection()
        self.apply_stored_formatting()
    
    def save_current_content(self):
        if not self.current_section or not self.current_book:
            return
        
        content = self.text_editor.get(1.0, tk.END).strip()
        book = self.books[self.current_book]
        book['content'][self.current_section] = content
        book['updated_at'] = datetime.now().isoformat()
        
        self.save_books()
        messagebox.showinfo("Saved", "Content saved successfully!")
    
    def update_word_count(self, event=None):
        content = self.text_editor.get(1.0, tk.END)
        words = len(content.split()) if content.strip() else 0
        self.word_count_label.config(text=f"Words: {words}")
    
    def export_book(self, book_id, format_type):
        book = self.books[book_id]
        
        if format_type == 'docx':
            try:
                from docx import Document
                from docx.shared import Inches
                
                doc = Document()
                
                # Title page
                title = doc.add_heading(book['title'], 0)
                title.alignment = 1  # Center
                
                author = doc.add_paragraph(f"by {book['author']}")
                author.alignment = 1
                
                doc.add_page_break()
                
                # Content
                for section_key, content in book.get('content', {}).items():
                    if content.strip():
                        section_title = section_key.replace('_', ' ').title()
                        doc.add_heading(section_title, 1)
                        
                        # Convert basic markdown
                        lines = content.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.startswith('# '):
                                doc.add_heading(line[2:], 1)
                            elif line.startswith('## '):
                                doc.add_heading(line[3:], 2)
                            elif line:
                                doc.add_paragraph(line)
                        
                        doc.add_page_break()
                
                # Save file
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx")],
                    initialname=f"{book['title']}.docx"
                )
                
                if filename:
                    doc.save(filename)
                    messagebox.showinfo("Success", f"Book exported to {filename}")
                    
            except ImportError:
                messagebox.showerror("Error", "python-docx not installed. Run: pip install python-docx")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {str(e)}")
    
    def delete_book(self, book_id):
        book = self.books[book_id]
        if messagebox.askyesno("Confirm Delete", f"Delete '{book['title']}'?\nThis cannot be undone."):
            del self.books[book_id]
            self.save_books()
            self.show_dashboard()
    
    def load_books(self):
        if self.books_file.exists():
            try:
                with open(self.books_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def save_books(self):
        with open(self.books_file, 'w', encoding='utf-8') as f:
            json.dump(self.books, f, indent=2, ensure_ascii=False)
    
    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = BooksyDesktop()
    app.run()